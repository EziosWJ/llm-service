"""文档解析器：将 txt/md/docx/pdf 文件解析为纯文本和章节结构，供下游分块使用。"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedSection:
    """解析后的文档章节，包含标题、层级和内容。"""
    heading: str
    level: int
    content: str


def parse_document(file_path: str | Path) -> dict[str, Any]:
    """根据文件扩展名分派到对应的解析函数，返回 {"text": ..., "sections": ...}。"""
    path = Path(file_path)
    ext = path.suffix.lower()
    logger.info("Parsing document: path=%s, ext=%s", path, ext)
    if ext in (".txt", ".md"):
        return _parse_txt(path)
    if ext == ".docx":
        return _parse_docx(path)
    if ext == ".pdf":
        return _parse_pdf(path)
    logger.warning("Unsupported document extension: %s", ext or "<none>")
    raise ValueError(f"Unsupported file type: {ext or '<none>'}")


def _parse_txt(path: Path) -> dict[str, Any]:
    """解析纯文本/md 文件，依次尝试 utf-8、gbk、latin-1 编码。"""
    last_err: UnicodeDecodeError | None = None
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            text = path.read_text(encoding=enc)
            return {"text": text, "sections": []}
        except UnicodeDecodeError as err:
            last_err = err
    if last_err:
        raise last_err
    return {"text": "", "sections": []}


def _parse_docx(path: Path) -> dict[str, Any]:
    """解析 docx 文件，提取全文和标题结构（Heading1/2/3）。"""
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx is required for .docx parsing") from exc

    document = docx.Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    full_text = "\n".join(paragraphs)

    sections: list[dict[str, Any]] = []
    for p in document.paragraphs:
        raw_text = p.text or ""
        text = raw_text.strip()
        if not text:
            continue
        style_name = getattr(getattr(p, "style", None), "name", "") or ""
        level = _extract_heading_level(style_name)
        if level is not None:
            sections.append({"heading": text, "level": level, "content": ""})

    return {"text": full_text, "sections": sections}


def _extract_heading_level(style_name: str) -> int | None:
    """从 Word 样式名中提取标题层级（1/2/3），非标题样式返回 None。"""
    lower = style_name.lower().replace(" ", "")
    if lower == "heading1":
        return 1
    if lower == "heading2":
        return 2
    if lower == "heading3":
        return 3
    return None


def _parse_pdf(path: Path) -> dict[str, Any]:
    """解析 PDF 文件，逐页提取文本并拼接。"""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required for .pdf parsing") from exc

    reader = PdfReader(str(path))
    texts: list[str] = []
    for page in reader.pages:
        content = page.extract_text() or ""
        if content.strip():
            texts.append(content.strip())
    return {"text": "\n".join(texts), "sections": []}
